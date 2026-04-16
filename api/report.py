from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import matplotlib
import matplotlib.pyplot as plt
import os
import re

# Ensure matplotlib runs headlessly
matplotlib.use('Agg')

# Design Tokens (Apple / Relayto inspired)
ACCENT_BLUE = "0284c7"    # Institutional Blue
ACCENT_GRAY = "64748b"    # Secondary
TEXT_MAIN = "0f172a"      # Obsidian
TEXT_SUB = "475569"       # Slate

def clean_relayto_text(text):
    """Refines LLM output for institutional reporting: strips markdown and control chars."""
    if not text: return "N/A"
    # Strip markdown bold/italic
    text = re.sub(r'\*\*|\*|_', '', text)
    # Strip bullet markers (keep the text)
    text = re.sub(r'^\s*[-•]\s*', '', text, flags=re.MULTILINE)
    # Clean control chars
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', str(text))

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for attr, val in edge_data.items():
                element.set(qn('w:{}'.format(attr)), str(val))

def add_institutional_rule(paragraph, color=ACCENT_BLUE, size="6"):
    """Adds a minimal horizontal rule above a section."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    top = pBdr.find(qn('w:top'))
    if top is None:
        top = OxmlElement('w:top')
        pBdr.append(top)
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), size)
    top.set(qn('w:color'), color)
    top.set(qn('w:space'), "12")

def create_minimalist_chart(score):
    plt.style.use('default')
    fig, ax = plt.subplots(figsize=(3, 2))
    color = '#0ea5e9' if score < 40 else '#f59e0b' if score < 70 else '#ef4444'
    
    # Clean bar
    ax.barh(["Risk Index"], [score], color=color, height=0.5)
    ax.barh(["Risk Index"], [100], color='#f1f5f9', height=0.5, zorder=0)
    
    ax.set_xlim(0, 100)
    ax.axis('off') # Hyper-minimalist Apple look
    
    ax.text(score, 0, f" {score}%", va='center', ha='left', fontsize=12, weight='bold', color=color)
    
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, transparent=True, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf

def parse_analysis_v2(text):
    """Robust institutional parser for the new high-narrative system prompt."""
    sections = {
        "summary": "N/A", "issues": "No identifying risks.", "explanation": "Pending audit.",
        "evidence": "No trace data.", "verdict": "Posturing.", "suggestions": "Review standard DISCO."
    }
    parts = {
        "summary": "Claim Summary:", "issues": "Key Issues:", "explanation": "Explanation:",
        "evidence": "Evidence & Proof:", "verdict": "Overall Verdict:", "suggestions": "Suggestions for Institutional Improvement:"
    }
    content_list = list(parts.items())
    for i in range(len(content_list)):
        key, header = content_list[i]
        if header in text:
            start = text.find(header) + len(header)
            end = len(text)
            if i + 1 < len(content_list):
                next_header = content_list[i+1][1]
                if next_header in text: end = text.find(next_header)
            sections[key] = clean_relayto_text(text[start:end].strip())
    return sections

def add_cover(doc):
    for _ in range(5): doc.add_paragraph()
    # Assets path fixed for Vercel deployment
    logo_path = os.path.join(os.path.dirname(__file__), "assets", "logo.png")
    if os.path.exists(logo_path):
         para = doc.add_paragraph()
         para.alignment = WD_ALIGN_PARAGRAPH.LEFT
         para.add_run().add_picture(logo_path, width=Inches(1.5))

    title = doc.add_paragraph()
    run = title.add_run("GreenLedger Forensic Intelligence")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEXT_MAIN)
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("INSTITUTIONAL AUDIT & RISK SYNTHESIS")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(ACCENT_BLUE)
    
    for _ in range(10): doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cells = table.rows[0].cells
    cells[0].text = "REPORT DATE\n" + datetime.now().strftime("%B %Y")
    cells[1].text = "REFERENCE ID\nGL-FORENSIC-2026-X"
    for cell in cells:
        for p in cell.paragraphs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor.from_string(TEXT_SUB)
    doc.add_page_break()

def add_executive_summary_v2(doc, claims_data):
    h = doc.add_heading("Executive Overview", level=1)
    add_institutional_rule(h)
    
    intro = doc.add_paragraph("This forensics report synthesizes institutional telemetry across the specified portfolio claims. Our objective is to evaluate the alignment between corporate disclosures and verified external trajectories.")
    intro.paragraph_format.space_after = Pt(24)

    # Clean Apple-Style table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    headers = ['ENTITY CLAIM', 'RISK POSTURE', 'VERDICT']
    for i, h_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h_text)
        run.font.size = Pt(8)
        run.font.bold = True
        set_cell_border(hdr_cells[i], bottom={"sz": 6, "color": "000000", "val": "single"})

    for item in claims_data:
        claim = item.get('claim', 'Unknown')
        analysis = item.get('analysis', {})
        row = table.add_row().cells
        row[0].text = claim[:50] + "..." if len(claim) > 50 else claim
        row[1].text = str(analysis.get("risk_level", "TBD"))
        row[2].text = "Institutional Gaps detected" if analysis.get("risk_score", 0) > 50 else "High Consistency"
        
        for cell in row:
            for p in cell.paragraphs: p.runs[0].font.size = Pt(9)
            set_cell_border(cell, bottom={"sz": 2, "color": "DDDDDD", "val": "single"})

    doc.add_page_break()

def add_claim_section(doc, item):
    claim = item.get('claim', 'Unknown Claim')
    analysis = item.get('analysis', {})
    sections = parse_analysis_v2(analysis.get("analysis", ""))
    risk_score = analysis.get("risk_score", 50)

    h1 = doc.add_heading(claim[:45] + "...", level=1)
    add_institutional_rule(h1)
    
    # Summary Highlight
    p = doc.add_paragraph()
    run = p.add_run(sections['summary'])
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_BLUE)
    p.paragraph_format.space_after = Pt(18)

    # Narrative Columns (Simulated with table)
    grid = doc.add_table(rows=1, cols=2)
    grid.columns[0].width = Inches(2)
    grid.columns[1].width = Inches(4.5)
    
    left = grid.cell(0, 0).paragraphs[0]
    left.add_run("FORENSIC INDEX\n").bold = True
    chart = create_minimalist_chart(risk_score)
    left.add_run().add_picture(chart, width=Inches(1.8))
    
    right = grid.cell(0, 1).paragraphs[0]
    right.add_run("Forensic Narrative\n").font.size = Pt(11)
    right.add_run(sections['explanation']).font.size = Pt(10)

    # Verdict Box
    doc.add_paragraph()
    v_head = doc.add_heading("Final Verdict", level=2)
    doc.add_paragraph(sections['verdict'])

    # Gaps & Institutional recommendations
    doc.add_heading("Identified Institutional Gaps", level=2)
    doc.add_paragraph(sections['issues'])

    doc.add_heading("Strategic Recommendations", level=2)
    doc.add_paragraph(sections['suggestions'])

    doc.add_page_break()

def generate_word_report(claims_data: list):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    add_cover(doc)
    add_executive_summary_v2(doc, claims_data)
    for item in claims_data:
        add_claim_section(doc, item)
        
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
